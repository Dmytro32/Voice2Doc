from docx import Document
import json
import os
def save(text):
   f = open("setting.json")

   data = json.load(f)
   file=data["doc"]["folder"]+data["doc"]["name"]
   if os.path.isfile(file):
      document = Document(file)
   else:
      document=Document()
   res=text.split('\n')
   for t in res:
      document.add_paragraph(t)
   document.save(file)




